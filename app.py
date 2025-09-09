import streamlit as st
from meal_generator import generate_meal_plan
import re
from docx import Document
from io import BytesIO
import base64
from datetime import date

def set_bg_from_local(img_path: str):
    with open(img_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    st.markdown(f"""
    <style>
    .stApp {{
      background: url("data:image/png;base64,{b64}") no-repeat center center fixed;
      background-size: cover;
    }}
    [data-testid="stHeader"] {{background: rgba(0,0,0,0);}}
    
    </style>
    """, unsafe_allow_html=True)

set_bg_from_local("fondo.png")  # pon aquí tu ruta

st.markdown("""
<style>
/* =========================
   PALETA (ajústala aquí)
   ========================= */
:root{
  --card:           #64D6ED;      /* fondo del bloque central */
  --card-weak:      #F0B34B;      /* paneles internos/headers */
  --card-strong:    #D4870A;      /* bordes/sombras sutiles */
  --surface:        #FFFFFF;      /* superficies de inputs */
  --surface-weak:   rgba(255,255,255,.85);
  --opciSupCol:     #457AFF;      /* fondo de las opciones superiores */
  --ink:            #1F2937;      /* texto principal */
  --ink-soft:       rgba(31,41,55,.85);
  --h1Color:        #FFB545;      /* título principal */
  --accent:         #1f7a8c;      /* títulos/enlaces/acciones */
  --accent-weak:    rgba(31,122,140,.12);
}

/* =========================
   FONDO GLOBAL + BLOQUE
   ========================= */
/* tu imagen de fondo ya la pones con set_bg_from_local() */
.stApp { color: var(--ink); }

/* ① BLOQUE CENTRAL (títulos, uploaders, formularios…)  */
.stMainBlockContainer{
  background: var(--card);
  padding: 2rem 2.4rem;
  border-radius: 20px;
  border: 1px solid color-mix(in srgb, var(--card-strong), #000 10%);
  box-shadow: 0 10px 30px rgba(0,0,0,.18);
}

/* =========================
   CABECERAS / TÍTULOS
   ========================= */
/* <div class="stHeading">…  h1 / h3  … */
.stHeading h1, .stHeading h2, .stHeading h3{
  color: var(--ink);
  letter-spacing:.2px;
}
#nutri-gen-generador-nutricional-interactivo{   /* tu H1 del título */
  background: var(--h1Color);
}

/* =========================
   FILE UPLOADERS
   ========================= */
/* contenedor del uploader (caja completa) */
[data-testid="stFileUploader"]{
  background: color-mix(in srgb, var(--card), #fff 10%);
  border-radius: 14px;
  border: 1px dashed color-mix(in srgb, var(--card-strong), #000 20%);
  padding:.35rem .35rem 0 .35rem;
}
/* zona de drop */
[data-testid="stFileUploaderDropzone"]{
  background: var(--surface-weak);
  border: 1px dashed rgba(0,0,0,.18);
  border-radius: 12px;
}
/* botón “Browse files” */
button[kind="secondary"]{
  background: var(--accent);
  color:#fff;
  border:0;
}
button[kind="secondary"]:hover{ filter:brightness(.95); }

/* zona de las opciones superiores */
[data-testid="stStatusWidget"]{
  background: var(--opciSupCol);
  border: 1px dashed rgba(0,0,0,.18);
  border-radius: 12px;
}

/* =========================
   FORM / CAMPOS
   ========================= */
/* evita que el form pinte su propia tarjeta blanca */
[data-testid="stForm"] > div{
  background: transparent !important;
  box-shadow:none !important;
  border:none !important;
}

/* inputs y selects */
[data-baseweb="input"] input,
textarea,
[data-baseweb="select"] > div{
  background: var(--surface);
  color: var(--ink);
}

/* sub-secciones con encabezado (las cajas que envuelven cada grupo) */
.stElementContainer .stHeading + div{
  background: color-mix(in srgb, var(--card), #fff 6%);
  border-radius: 12px;
}

/* =========================
   HEADERS DE SECCIÓN (H3)
   ========================= */
/* Ej: #diagnostico-nutricional, #objetivos-del-plan-nutricional,
      #estrategia-nutricional, #reparto-de-macronutrientes-diario-aproximado */
h3#datos-personales,
h3#diagnostico-nutricional,
h3#objetivos-del-plan-nutricional,
h3#estrategia-nutricional,
h3#reparto-de-macronutrientes-diario-aproximado,
h3#recomendaciones-generales,
h3#indicaciones-y-sugerencias-personalizadas,
h3#seguimiento-y-reevaluacion{
  padding:.6rem .75rem;
  margin:.75rem 0 .5rem 0;
  background: var(--card-weak);
  border-radius:10px;
  color: var(--ink);
}

/* =========================
   ALERT/INFO (el recuadro gris al final)
   ========================= */
[data-testid="stAlertContainer"]{
  background: var(--accent-weak) !important;
  border: 1px solid color-mix(in srgb, var(--accent), #000 15%);
  color: var(--ink-soft);
  border-radius: 12px;
}

/* =========================
   BOTÓN “Generar Plan”
   ========================= */
.stButton button{
  background: var(--accent);
  color:#fff;
  border:0;
  border-radius:10px;
  padding:.5rem 1rem;
  box-shadow: 0 4px 10px rgba(0,0,0,.15);
}
.stButton button:hover{ filter:brightness(.95); }

/* =========================
   DETALLES
   ========================= */
hr{ border-color: color-mix(in srgb, var(--card-strong), #000 10%); }
a{ color: var(--accent); }
</style>
""", unsafe_allow_html=True)

# =====================
# Helpers numéricos
# =====================
def _to_float(x):
    if x is None:
        return None
    x = str(x).strip().replace(",", ".")
    try:
        return float(x)
    except:
        return None

# =====================
# Parser de la FICHA (Plan A.docx)
# =====================
def parse_docx_plan_a(file) -> dict:
    """
    Extrae datos desde el .docx de ficha (Plan A).
    Busca encabezados y captura el texto siguiente hasta el próximo encabezado numerado.
    """
    doc = Document(file)
    text = "\n".join(p.text for p in doc.paragraphs).replace("\xa0", " ")

    # Línea "Etiqueta: valor" (tolerante a acentos, espacios y variaciones mínimas)
    def get_line(label_regex: str):
        # ^...: valor  (hasta fin de línea)
        pattern = rf"(?mi)^\s*{label_regex}\s*:\s*(?P<val>.+?)\s*$"
        m = re.search(pattern, text)
        return m.group("val").strip() if m else None

    # Bloques entre encabezados numerados "N. Título" hasta el siguiente "M. ..."
    def get_block(title_regex: str):
        # Busca “N. <título> … <contenido> … (hasta la próxima línea que empiece con número y punto)”
        # El DOTALL se consigue con (?s) dentro del patrón
        pattern = rf"(?ims)^\s*{title_regex}\s*\n(?P<body>.*?)(?=^\s*\d+\.\s|\Z)"
        m = re.search(pattern, text)
        return (m.group("body").strip() if m else "")

    # Dentro de un bloque, “Subetiqueta: valor” (captura solo esa línea)
    def get_subline(block_text: str, sublabel_regex: str):
        m = re.search(rf"(?mi)^\s*{sublabel_regex}\s*:\s*(?P<val>.+?)\s*$", block_text)
        return m.group("val").strip() if m else ""

    # ---------------- 1. DATOS DEL CLIENTE ----------------
    datos = {
        "nombre":          get_line(r"Nombre completo"),
        "edad":            _to_float(get_line(r"Edad")),
        "sexo":            get_line(r"Sexo"),
        "peso":            _to_float(get_line(r"Peso\s*\(kg\)")),
        "estatura":        _to_float(get_line(r"Estatura\s*\(cm\)")),
        "imc":             _to_float(get_line(r"IMC")),
        "grasa":           _to_float(get_line(r"%\s*Grasa\s*corporal")),
        "masa_muscular":   _to_float(get_line(r"%\s*Masa\s*muscular")),
        "objetivo":        get_line(r"Objetivo\s+nutricional\s+principal"),
    }

    # ---------------- 2. DIAGNÓSTICO NUTRICIONAL ----------------
    diag = get_block(r"2\.\s*Diagn[óo]stico\s+Nutricional")
    datos["desc_est_nutricional"] = get_subline(diag, r"Descripci[óo]n del estado nutricional actual") or diag
    datos["res_analisis"]        = get_subline(diag, r"Resultados de an[áa]lisis relevantes\s*\(si aplica\)") or ""
    datos["obs_clinicas"]        = get_subline(diag, r"Observaciones cl[ií]nicas y de h[áa]bitos") or ""

    # ---------------- 3. OBJETIVOS DEL PLAN ----------------
    objetivos = get_block(r"3\.\s*Objetivos del Plan Nutricional")
    datos["objetivo_principal"]   = get_subline(objetivos, r"Objetivo principal\s*\(p[ée]rdida de peso, ganancia muscular, etc\.\)") or (datos["objetivo"] or "")
    datos["objetivos_secundarios"]= get_subline(objetivos, r"Objetivos secundarios") or ""
    datos["plazo_estimado"]       = get_subline(objetivos, r"Plazo estimado para alcanzar los objetivos") or ""

    # ---------------- 4. ESTRATEGIA NUTRICIONAL ----------------
    estrategia = get_block(r"4\.\s*Estrategia Nutricional")
    datos["alimentacion_recomendada"] = get_subline(estrategia, r"Tipo de alimentaci[óo]n recomendada\s*\(ej\.\s*mediterr[áa]nea, hipocal[óo]rica, cetog[ée]nica, etc\.\)")
    datos["just_plan"]                = get_subline(estrategia, r"Justificaci[óo]n de la elecci[óo]n del plan")
    # El doc pone “Restricciones o preferencias alimentarias consideradas:”
    datos["restricc_pref"]           = get_subline(estrategia, r"Restricciones o preferencias alimentarias consideradas")

    # ---------------- 5. REPARTO DE MACROS (si viniera ya en la ficha) ----------------
    # En tu texto aparecen “Calorias totales (g)”, etc.  Aceptamos (g) opcional.
    datos["calorias"]      = _to_float(get_line(r"Calor[ií]as?\s+totales(?:\s*\(g\))?"))
    datos["proteinas"]     = _to_float(get_line(r"Prote[ií]nas?(?:\s*\(g\))?"))
    datos["grasas"]        = _to_float(get_line(r"Grasas?(?:\s*\(g\))?"))
    datos["azucares"]      = _to_float(get_line(r"Az[úu]cares?(?:\s*\(g\))?"))
    datos["carbohidratos"] = _to_float(get_line(r"Carbohidratos?(?:\s*\(g\))?"))

    # ---------------- 6. RECOMENDACIONES GENERALES ----------------
    rec = get_block(r"6\.\s*Recomendaciones Generales")
    datos["hidratacion_litros"] = _to_float(get_subline(rec, r"Hidrataci[óo]n:\s*m[ií]nimo\s+(\d+(?:[.,]\d+)?)\s*litros/d[ií]a") or get_subline(rec, r"Hidrataci[óo]n"))
    datos["freq_act_fis"]       = get_subline(rec, r"Frecuencia de actividad f[ií]sica sugerida")
    datos["cal_horas_sueno"]    = get_subline(rec, r"Calidad y horas de sue[ñn]o recomendadas")
    datos["suplementacion"]     = get_subline(rec, r"Suplementaci[óo]n recomendada\s*\(si aplica\)")

    # ---------------- 7. INDICACIONES PERSONALIZADAS ----------------
    indic = get_block(r"7\.\s*Indicaciones y Sugerencias Personalizadas")
    datos["consejos_especificos"] = indic.strip()

    # ---------------- 8. SEGUIMIENTO Y REEVALUACIÓN ----------------
    seg = get_block(r"8\.\s*Seguimiento y Reevaluaci[óo]n")
    datos["prim_rev_sugerida"] = get_subline(seg, r"Primera revisi[óo]n sugerida en")
    datos["param_eval_seg"]    = get_subline(seg, r"Par[áa]metros a evaluar en seguimiento")
    datos["frec_rec_rev"]      = get_subline(seg, r"Frecuencia recomendada de revisiones")

    # ---------------- 10. OBSERVACIONES FINALES ----------------
    obs = get_block(r"10\.\s*Observaciones Finales del Profesional")
    datos["observaciones_finales"] = obs.strip()

    return datos

# =====================
# Splitter de secciones del plan
# =====================
def split_plan_sections(plan_text: str) -> dict:
    sections = {"DESAYUNO": "", "COMIDA": "", "CENA": "", "MERIENDA": ""}

    header_re = re.compile(
        r"""(?im)
        ^\s*
        (?:[#>*-]+\s*)?
        (?:\d+\.\s*)?
        (?:\*\*|__)?\s*
        (desayunos?|comidas?|almuerzos?|cenas?|meriendas?|snacks?|merienda\s*o\s*snack)
        \s*(?:\*\*|__)?\s*
        :?\s*$
        """,
        re.VERBOSE,
    )

    matches = list(header_re.finditer(plan_text))
    if not matches:
        return sections

    bounds = [(m.group(1).lower(), m.start(), m.end()) for m in matches]
    bounds.append(("__end__", len(plan_text), len(plan_text)))

    for i in range(len(bounds) - 1):
        name, _start, end = bounds[i]
        next_start = bounds[i + 1][1]
        body = plan_text[end:next_start].strip()

        if name.startswith("desayuno"):
            key = "DESAYUNO"
        elif name.startswith(("comida", "almuerzo")):
            key = "COMIDA"
        elif name.startswith("cena"):
            key = "CENA"
        elif name.startswith(("merienda", "snack", "merienda o snack")):
            key = "MERIENDA"
        else:
            continue
        sections[key] = (sections[key] + "\n" + body).strip() if sections[key] else body

    return sections

# =====================
# Relleno DOCX preservando estilo del placeholder
# =====================
def fill_docx_template(template_file, mapping: dict) -> BytesIO:
    doc = Document(template_file)

    def copy_format(src_run, dst_run):
        dst_run.bold = src_run.bold
        dst_run.italic = src_run.italic
        dst_run.underline = src_run.underline
        if src_run.font is not None:
            dst_run.font.name = src_run.font.name
            dst_run.font.size = src_run.font.size
            dst_run.font.color.rgb = getattr(src_run.font.color, "rgb", None)

    def add_text_with_breaks(run, text):
        parts = str(text).split("\n")
        run.text = parts[0]
        for part in parts[1:]:
            run.add_break()
            run.add_text(part)

    def replace_in_paragraph(p):
        if not p.runs:
            return
        full = "".join(r.text for r in p.runs)
        if not full:
            return
        run_map = []
        for idx, r in enumerate(p.runs):
            run_map += [idx] * len(r.text)

        segments = []
        i = 0
        while i < len(full):
            hit_key = None
            for k in mapping.keys():
                if full.startswith(k, i):
                    if hit_key is None or len(k) > len(hit_key):
                        hit_key = k
            if hit_key:
                start_run_idx = run_map[i] if i < len(run_map) else 0
                segments.append(("__REPL__", mapping[hit_key], start_run_idx))
                i += len(hit_key)
                continue
            start_i = i
            start_run_idx = run_map[i] if i < len(run_map) else 0
            while i < len(full):
                next_is_key = any(full.startswith(k, i) for k in mapping.keys())
                if next_is_key:
                    break
                curr_run_idx = run_map[i] if i < len(run_map) else start_run_idx
                if curr_run_idx != start_run_idx:
                    break
                i += 1
            segments.append(("__TEXT__", full[start_i:i], start_run_idx))

        for r in p.runs:
            r.text = ""
        base_run = p.runs[0] if p.runs else p.add_run()

        first = True
        for kind, content, src_idx in segments:
            src_run = p.runs[src_idx] if src_idx < len(p.runs) else base_run
            tgt = base_run if first else p.add_run()
            first = False
            copy_format(src_run, tgt)
            add_text_with_breaks(tgt, content if kind == "__REPL__" else content)

    def replace_in_table(t):
        for row in t.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    replace_in_paragraph(cp)
                for subt in cell.tables:
                    replace_in_table(subt)

    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for t in doc.tables:
        replace_in_table(t)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

def build_basic_docx(
    nombre:str,
    user_data:dict,
    secs:dict,
    desc_act_nutricional_state:str,
    relevant_analy_res:str,
    clin_hab_observs:str,
    objetivos_secundarios:str,
    plazo_estimado:str,
    al_rec:str,
    just_plan:str,
    restrictions:str,
    hidra,
    fr_act_fis:str,
    cal_horas_sueno:str,
    supl_rec:str,
    cons_esp:str,
    prim_rev_sug:str,
    param_ev_seg:str,
    frec_rec_rev:str,
    reparto_macros_text:str
) -> BytesIO:
    """Crea un DOCX autónomo (sin plantilla) con todo el contenido del plan."""
    d = Document()
    d.add_heading(f"Plan nutricional personalizado para {nombre or 'Cliente'}", 0)
    d.add_paragraph(f"Fecha: {date.today().strftime('%d/%m/%Y')}")

    d.add_heading("1. Datos del Cliente", level=1)
    p = d.add_paragraph()
    p.add_run("Nombre completo: ").bold = True; d.add_paragraph(str(user_data.get("nombre") or "—"))
    p = d.add_paragraph(); p.add_run("Edad: ").bold = True; d.add_paragraph(str(user_data.get("edad") or "—"))
    p = d.add_paragraph(); p.add_run("Sexo: ").bold = True; d.add_paragraph(str(user_data.get("sexo") or "—"))
    p = d.add_paragraph(); p.add_run("Peso (kg): ").bold = True; d.add_paragraph(str(user_data.get("peso") or "—"))
    p = d.add_paragraph(); p.add_run("Estatura (cm): ").bold = True; d.add_paragraph(str(user_data.get("estatura") or "—"))
    p = d.add_paragraph(); p.add_run("IMC: ").bold = True; d.add_paragraph(str(user_data.get("imc") or "—"))
    p = d.add_paragraph(); p.add_run("% Grasa corporal: ").bold = True; d.add_paragraph(str(user_data.get("grasa") or "—"))
    p = d.add_paragraph(); p.add_run("% Masa muscular: ").bold = True; d.add_paragraph(str(user_data.get("masa_muscular") or "—"))
    p = d.add_paragraph(); p.add_run("Objetivo nutricional principal: ").bold = True; d.add_paragraph(str(user_data.get("objetivo") or "—"))

    d.add_heading("2. Diagnóstico Nutricional", level=1)
    d.add_paragraph("Descripción del estado nutricional actual:", style=None).runs[0].bold = True
    d.add_paragraph(desc_act_nutricional_state or "—")
    d.add_paragraph("Resultados de análisis relevantes (si aplica):").runs[0].bold = True
    d.add_paragraph(relevant_analy_res or "—")
    d.add_paragraph("Observaciones clínicas y de hábitos:").runs[0].bold = True
    d.add_paragraph(clin_hab_observs or "—")

    d.add_heading("3. Objetivos del Plan Nutricional", level=1)
    d.add_paragraph(f"Objetivo principal: {user_data.get('objetivo') or '—'}")
    d.add_paragraph(f"Objetivos secundarios: {objetivos_secundarios or '—'}")
    d.add_paragraph(f"Plazo estimado: {plazo_estimado or '—'}")

    d.add_heading("4. Estrategia Nutricional", level=1)
    d.add_paragraph(f"Tipo de alimentación recomendada: {al_rec or '—'}")
    d.add_paragraph("Justificación de la elección del plan:").runs[0].bold = True
    d.add_paragraph(just_plan or "—")
    d.add_paragraph(f"Restricciones o preferencias alimentarias consideradas: {restrictions or '—'}")

    d.add_heading("5. Reparto de Macronutrientes (Diario Aproximado)", level=1)
    for line in (reparto_macros_text or "").splitlines():
        d.add_paragraph(line)

    d.add_heading("6. Distribución de Comidas (Ejemplo Diario)", level=1)
    d.add_paragraph("DESAYUNOS").runs[0].bold = True
    d.add_paragraph(secs.get("DESAYUNO") or "—")
    d.add_paragraph("COMIDAS").runs[0].bold = True
    d.add_paragraph(secs.get("COMIDA") or "—")
    d.add_paragraph("CENAS").runs[0].bold = True
    d.add_paragraph(secs.get("CENA") or "—")
    d.add_paragraph("MERIENDAS / SNACKS").runs[0].bold = True
    d.add_paragraph(secs.get("MERIENDA") or "—")

    d.add_heading("7. Recomendaciones Generales", level=1)
    d.add_paragraph(f"Hidratación: mínimo {hidra} litros/día")
    d.add_paragraph(f"Frecuencia de actividad física sugerida: {fr_act_fis or '—'}")
    d.add_paragraph(f"Calidad y horas de sueño recomendadas: {cal_horas_sueno or '—'}")
    d.add_paragraph(f"Suplementación recomendada (si aplica): {supl_rec or '—'}")

    d.add_heading("8. Indicaciones y Sugerencias Personalizadas", level=1)
    d.add_paragraph(cons_esp or "—")

    d.add_heading("9. Seguimiento y Reevaluación", level=1)
    d.add_paragraph(f"Primera revisión sugerida en: {prim_rev_sug or '—'}")
    d.add_paragraph(f"Parámetros a evaluar en seguimiento: {param_ev_seg or '—'}")
    d.add_paragraph(f"Frecuencia recomendada de revisiones: {frec_rec_rev or '—'}")

    # Guardar a memoria
    buf = BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


def build_basic_pdf(
    nombre:str,
    user_data:dict,
    secs:dict,
    texto_bloques:list
) -> BytesIO:
    """
    Crea un PDF sencillo (A4) usando reportlab.
    texto_bloques: lista de (titulo, [lineas]) para maquetar bloques.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
        from reportlab.lib.utils import simpleSplit
    except Exception:
        return None  # reportlab no está disponible

    W, H = A4
    margin = 2 * cm
    y = H - margin

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setTitle(f"Plan nutricional - {nombre or 'Cliente'}")

    def draw_title(txt, size=16):
        nonlocal y
        if y < 3*cm:
            c.showPage(); y = H - margin
        c.setFont("Helvetica-Bold", size)
        c.drawString(margin, y, txt)
        y -= 0.8*cm

    def draw_par(lines, size=11, leading=14):
        nonlocal y
        c.setFont("Helvetica", size)
        for line in lines:
            parts = simpleSplit(str(line), "Helvetica", size, W - 2*margin)
            for p in parts:
                if y < 2.5*cm:
                    c.showPage(); y = H - margin; c.setFont("Helvetica", size)
                c.drawString(margin, y, p)
                y -= leading

    # Cabecera
    draw_title(f"Plan nutricional personalizado para {nombre or 'Cliente'}", 18)
    draw_par([f"Fecha: {date.today().strftime('%d/%m/%Y')}"])

    # Datos cortos en una lista
    draw_title("1. Datos del Cliente", 14)
    datos_lines = [
        f"Edad: {user_data.get('edad','—')}",
        f"Sexo: {user_data.get('sexo','—')}",
        f"Peso (kg): {user_data.get('peso','—')}",
        f"Estatura (cm): {user_data.get('estatura','—')}",
        f"IMC: {user_data.get('imc','—')}",
        f"% Grasa: {user_data.get('grasa','—')}",
        f"% Masa muscular: {user_data.get('masa_muscular','—')}",
        f"Objetivo: {user_data.get('objetivo','—')}",
    ]
    draw_par(datos_lines)

    # Bloques con títulos
    for titulo, lineas in texto_bloques:
        draw_title(titulo, 14)
        draw_par(lineas)

    # Comidas
    draw_title("6. Distribución de Comidas (Ejemplo Diario)", 14)
    for k in ("DESAYUNO", "COMIDA", "CENA", "MERIENDA"):
        draw_par([k + ":"], size=12)
        draw_par([(secs.get(k) or "—")])

    c.showPage()
    c.save()
    buf.seek(0)
    return buf


# =====================
# UI
# =====================
st.title("🥗 NutriGen - Generador Nutricional Interactivo 🥗")

uploaded = st.file_uploader("Sube la FICHA del cliente (Plan A .docx)", type=["docx"])
template_file = st.file_uploader("Sube la PLANTILLA a rellenar (.docx)", type=["docx"])

parsed = {}
if uploaded is not None:
    try:
        parsed = parse_docx_plan_a(uploaded) or {}
        with st.expander("📄 Datos importados del Plan A"):
            st.write(parsed)
    except Exception as e:
        st.error(f"No se pudo leer el documento: {e}")

with st.form("nutri_form"):
    st.header("Introduce / revisa los datos")

    st.subheader("Datos personales")
    nombre = st.text_input("Nombre completo", value=parsed.get("nombre", "") or "")
    edad = st.number_input("Edad", 0, 120, int(parsed.get("edad") or 0))
    sexo = st.selectbox("Sexo", ["Hombre", "Mujer", "Otro"],
                        index={"hombre":0,"mujer":1}.get(str(parsed.get("sexo") or "").lower(), 2))

    weight = st.number_input("Peso (kg)", 0.0, 250.0, float(parsed.get("peso") or 0.0))
    height = st.number_input("Estatura (cm)", 0.0, 250.0, float(parsed.get("estatura") or 0.0))
    imc_val = parsed.get("imc")
    st.markdown(f"**IMC (auto):** {imc_val if imc_val is not None else '—'}")

    grasa = st.number_input("% Grasa corporal", 0.0, 100.0, float(parsed.get("grasa") or 0.0))
    masa = st.number_input("% Masa muscular", 0.0, 100.0, float(parsed.get("masa_muscular") or 0.0))
    objetivo_principal = st.text_input("Objetivo nutricional principal", value=parsed.get("objetivo","") or "")

    st.markdown("---")
    st.subheader("Diagnóstico Nutricional")
    desc_act_nutricional_state = st.text_area("Descripción del estado nutricional actual", value=parsed.get("desc_est_nutricional",""))
    relevant_analy_res = st.text_area("Resultados de análisis relevantes (si aplica)", value=parsed.get("res_analisis",""))
    clin_hab_observs = st.text_area("Observaciones clínicas y de hábitos", value=parsed.get("obs_clinicas",""))

    st.markdown("---")
    st.subheader("Objetivos del Plan Nutricional")
    objetivos_secundarios = st.text_input("Objetivos secundarios", value=parsed.get("objetivos_secundarios",""))
    plazo_estimado = st.text_input("Plazo estimado para alcanzar los objetivos", value=parsed.get("plazo_estimado",""))

    st.markdown("---")
    st.subheader("Estrategia Nutricional")
    al_rec = st.text_input("Tipo de alimentación recomendada (ej. mediterránea, hipocalórica, cetogénica, etc.)", 
                           value=parsed.get("alimentacion_recomendada",""))
    just_plan = st.text_area("Justificación de la elección del plan", value=parsed.get("just_plan",""))
    restrictions = st.text_input("Restricciones o preferencias alimentarias consideradas", 
                                 value=parsed.get("restricc_pref",""))
    # En la ficha no hay “lista de alimentos”, así que lo mantengo aparte:
    preferences = st.text_input("Preferencias alimentarias (opcional para el generador)")
    menu_input = st.text_area("Lista de alimentos disponibles (opcional para el generador)")

    st.markdown("---")
    st.subheader("Reparto de Macronutrientes (Diario Aproximado)")
    calories      = st.number_input("Calorías diarias objetivo", 1000, 5000, int(parsed.get("calorias") or 2000))
    protein       = st.number_input("Proteína diaria objetivo (g)", 0, 300,  int(parsed.get("proteinas") or 120))
    fat           = st.number_input("Grasa diaria objetivo (g)", 0, 300,     int(parsed.get("grasas") or 60))
    sugar         = st.number_input("Azúcar diaria objetivo (g)", 0, 300,    int(parsed.get("azucares") or 20))
    carbohydrates = st.number_input("Carbohidratos diarios objetivo (g)", 0, 500, int(parsed.get("carbohidratos") or 210))

    st.markdown("---")
    st.subheader("Recomendaciones Generales")
    # La ficha dice: "Hidratación: mínimo __ litros/día"
    hidra        = st.number_input("Hidratación (litros/día)", 0.0, 6.0, float(parsed.get("hidratacion_litros") or 2.0), step=0.1)
    fr_act_fis   = st.text_input("Frecuencia de actividad física sugerida", value=parsed.get("freq_act_fis",""))
    cal_horas_sueno = st.text_input("Calidad y horas de sueño recomendadas", value=parsed.get("cal_horas_sueno",""))
    supl_rec     = st.text_input("Suplementación recomendada (si aplica)", value=parsed.get("suplementacion",""))

    st.markdown("---")
    st.subheader("Indicaciones y Sugerencias Personalizadas")
    cons_esp = st.text_area("Consejos específicos adaptados al estilo de vida del cliente, hábitos y horarios", 
                            value=parsed.get("consejos_especificos",""))

    st.markdown("---")
    st.subheader("Seguimiento y Reevaluación")
    prim_rev_sug = st.text_input("Primera revisión sugerida en", value=parsed.get("prim_rev_sugerida",""))
    param_ev_seg = st.text_input("Parámetros a evaluar en seguimiento", value=parsed.get("param_eval_seg",""))
    frec_rec_rev = st.text_input("Frecuencia recomendada de revisiones", value=parsed.get("frec_rec_rev",""))

    submitted = st.form_submit_button("Generar Plan")

# =====================
# Generar plan + guardar en sesión
# =====================
if submitted:
    user_data = {
        "nombre": nombre,
        "edad": edad,
        "sexo": sexo,
        "peso": weight,
        "estatura": height,
        "imc": imc_val,
        "grasa": grasa,
        "masa_muscular": masa,
        "objetivo": objetivo_principal,
        "preferences": preferences,
        "restrictions": restrictions,
        "calories": calories,
        "protein": protein,
        "sugar": sugar,
    }

    st.markdown("### 🧠 Resultado:")
    with st.spinner("Generando el plan..."):
        plan = generate_meal_plan(user_data, menu_input)
        st.session_state["user_data"] = user_data
        st.session_state["plan"] = plan
        st.session_state["has_plan"] = True
        secs = split_plan_sections(plan)
        st.session_state["plan_sections"] = secs

        st.markdown("---")
        blocks = re.split(r"\n-{3,}\n", plan) if "---" in plan else [plan]
        for block in blocks:
            st.markdown(block)

has_plan = st.session_state.get("has_plan", False)
saved_user_data = st.session_state.get("user_data")
saved_plan = st.session_state.get("plan")

# =====================
# Rellenar PLANTILLA y descargar
# =====================
# =====================
# Descargas (plantilla si hay + básico Word/PDF siempre que exista plan)
# =====================
if has_plan and saved_user_data and saved_plan:
    user_data = saved_user_data
    plan = saved_plan
    secs = st.session_state.get("plan_sections", {})

    # % macros (como ya hacías)
    try:
        pct_prot = round((user_data["protein"] * 4) / user_data["calories"] * 100)
    except Exception:
        pct_prot = 20
    pct_grasa = 30
    pct_carb = max(0, 100 - pct_prot - pct_grasa)

    reparto_macros_text = (
        f"Calorías totales: {user_data['calories']} kcal\n"
        f"Proteínas: ~{pct_prot}%\n"
        f"Grasas: ~{pct_grasa}%\n"
        f"Carbohidratos: ~{pct_carb}%"
    )

    # ---------- BOTONES “BÁSICO” (siempre disponibles) ----------
    # Construimos el DOCX básico
    basic_docx = build_basic_docx(
        nombre,
        user_data,
        secs,
        desc_act_nutricional_state,
        relevant_analy_res,
        clin_hab_observs,
        objetivos_secundarios,
        plazo_estimado,
        al_rec,
        just_plan,
        restrictions,
        hidra,
        fr_act_fis,
        cal_horas_sueno,
        supl_rec,
        cons_esp,
        prim_rev_sug,
        param_ev_seg,
        frec_rec_rev,
        reparto_macros_text
    )
    st.download_button(
        "⬇️ Descargar Word (básico)",
        data=basic_docx,
        file_name=f"Plan_{(nombre or 'cliente').replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    # Construimos el PDF básico (si reportlab está disponible)
    pdf_blocks = [
        ("2. Diagnóstico Nutricional", [
            f"Descripción: {desc_act_nutricional_state or '—'}",
            f"Análisis relevantes: {relevant_analy_res or '—'}",
            f"Observaciones de hábitos: {clin_hab_observs or '—'}",
        ]),
        ("3. Objetivos del Plan Nutricional", [
            f"Objetivo principal: {user_data.get('objetivo') or '—'}",
            f"Objetivos secundarios: {objetivos_secundarios or '—'}",
            f"Plazo estimado: {plazo_estimado or '—'}",
        ]),
        ("4. Estrategia Nutricional", [
            f"Alimentación recomendada: {al_rec or '—'}",
            "Justificación:",
            (just_plan or "—"),
            f"Restricciones/Preferencias: {restrictions or '—'}",
        ]),
        ("5. Reparto de Macronutrientes (Diario Aproximado)", reparto_macros_text.splitlines()),
        ("7. Recomendaciones Generales", [
            f"Hidratación: mínimo {hidra} L/día",
            f"Actividad física: {fr_act_fis or '—'}",
            f"Sueño: {cal_horas_sueno or '—'}",
            f"Suplementación: {supl_rec or '—'}",
        ]),
        ("8. Indicaciones y Sugerencias Personalizadas", [cons_esp or "—"]),
        ("9. Seguimiento y Reevaluación", [
            f"Primera revisión: {prim_rev_sug or '—'}",
            f"Parámetros a evaluar: {param_ev_seg or '—'}",
            f"Frecuencia de revisiones: {frec_rec_rev or '—'}",
        ]),
    ]
    basic_pdf = build_basic_pdf(nombre, user_data, secs, pdf_blocks)
    if basic_pdf is None:
        st.info("⚠️ Para exportar a PDF instala `reportlab` (pip install reportlab).")
    else:
        st.download_button(
            "⬇️ Descargar PDF (básico)",
            data=basic_pdf,
            file_name=f"Plan_{(nombre or 'cliente').replace(' ', '_')}.pdf",
            mime="application/pdf",
        )

    # ---------- SI HAY PLANTILLA, además ofrece tu descarga con placeholders ----------
    if template_file is not None:
        mapping = {
            "{{NOMBRE_COMPLETO}}": str(nombre or "—"),
            "{{EDAD}}": str(edad or "—"),
            "{{SEXO}}": str(sexo or "—"),
            "{{PESO_KG}}": f"{float(user_data.get('peso') or 0):.1f}" if user_data.get("peso") else "—",
            "{{ESTATURA_CM}}": f"{float(user_data.get('estatura') or 0):.1f}" if user_data.get("estatura") else "—",
            "{{IMC}}": f"{float(user_data.get('imc')):.1f}" if user_data.get("imc") is not None else "—",
            "{{GRASA_PCT}}": f"{float(user_data.get('grasa') or 0):.1f}" if user_data.get("grasa") is not None else "—",
            "{{MASA_MUSCULAR_PCT}}": f"{float(user_data.get('masa_muscular') or 0):.1f}" if user_data.get("masa_muscular") is not None else "—",
            "{{OBJETIVO_PRINCIPAL}}": str(user_data.get("objetivo") or "—"),

            "{{DESC_EST_NUT_ACT}}": str(desc_act_nutricional_state or "—"),
            "{{RES_AN_R}}": str(relevant_analy_res or "—"),
            "{{OBS_CLI_HAB}}": str(clin_hab_observs or "—"),

            "{{OBJS_SECS}}": str(objetivos_secundarios or "—"),
            "{{PLA_EST}}": str(plazo_estimado or "—"),

            "{{AL_REC}}": str(al_rec or "—"),
            "{{JUST_PL}}": str(just_plan or "—"),
            "{{REST_PREF_AL}}": str(restrictions or "—"),

            "{{CALORIAS_OBJ}}": str(calories),
            "{{PROTEINA_OBJ}}": str(protein),
            "{{AZUCAR_OBJ}}": str(sugar),
            "{{CARB_OBJ}}": str(carbohydrates),
            "{{GR _OBJ}}": str(fat),  # ese placeholder raro de tu plantilla

            "{{DESAYUNO}}": secs.get("DESAYUNO", "") or "—",
            "{{COMIDA}}": secs.get("COMIDA", "") or "—",
            "{{CENA}}": secs.get("CENA", "") or "—",
            "{{MERIENDA}}": secs.get("MERIENDA", "") or "—",

            "{{HIDRA}}": str(hidra),
            "{{FR_ACT_FIS}}": str(fr_act_fis or "—"),
            "{{CAL_HORS_SUEÑ_REC}}": str(cal_horas_sueno or "—"),
            "{{SUPL_REC}}": str(supl_rec or "—"),

            "{{CONS_ESP_EST_CLI}}": str(cons_esp or "—"),
            "{{PRIM_REV_SUG}}": str(prim_rev_sug or "—"),
            "{{PARAM_EV_SEG}}": str(param_ev_seg or "—"),
            "{{FREC_REC_REV}}": str(frec_rec_rev or "—"),

            "{{REPARTO_MACROS}}": reparto_macros_text,
            "{{RECOMENDACIONES}}": "- Hidratación: mantener 1.5–2 L/día.\n- Verduras en 2+ comidas.\n- Ajustar raciones a saciedad.",
            "{{INDICACIONES}}": "Priorizar alimentos frescos y proteína magra. Reducir ultraprocesados.",
            "{{SEGUIMIENTO}}": "Revisión en 4–6 semanas.",
            "{{OBSERVACIONES}}": parsed.get("observaciones_finales","—"),
        }

        filled = fill_docx_template(template_file, mapping)
        nombre_archivo = f"Plan_{(nombre or 'cliente').replace(' ', '_')}.docx"
        st.download_button(
            "⬇️ Descargar plan (plantilla rellenada)",
            data=filled,
            file_name=nombre_archivo,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
else:
    st.info("Genera el plan para habilitar las descargas.")

