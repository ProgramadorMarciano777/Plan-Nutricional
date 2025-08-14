import streamlit as st
from meal_generator import generate_meal_plan
import re
from docx import Document
from io import BytesIO
import re

# -------- Helpers --------
def _to_float(x):
    if x is None: 
        return None
    x = str(x).strip().replace(',', '.')
    try:
        return float(x)
    except:
        return None

def parse_docx(file) -> dict:
    """
    Espera líneas tipo:
    Nombre completo: ...
    Edad: 30
    Sexo: Hombre/Mujer
    Peso (kg): 78
    Estatura (cm): 175
    IMC: 25.5
    % Grasa corporal: 18
    % Masa muscular: 40
    Objetivo nutricional principal: Pérdida de grasa
    """
    doc = Document(file)
    text = "\n".join(p.text for p in doc.paragraphs)

    def get(label):
        # Busca "Label:" con may/min, acentos y espacios flexibles
        import re
        pattern = rf"{label}\s*:\s*(.+)"
        m = re.search(pattern, text, flags=re.IGNORECASE)
        return m.group(1).strip() if m else None

    data = {
        "nombre": get(r"Nombre completo"),
        "edad": _to_float(get(r"Edad")),
        "sexo": get(r"Sexo"),
        "peso": _to_float(get(r"Peso\s*\(kg\)")),
        "estatura": _to_float(get(r"Estatura\s*\(cm\)")),
        "imc": _to_float(get(r"IMC")),
        "grasa": _to_float(get(r"%\s*Grasa\s*corporal")),
        "masa_muscular": _to_float(get(r"%\s*Masa\s*muscular")),
        "objetivo": get(r"Objetivo\s*nutricional\s*principal"),
    }

    # Si falta IMC y tenemos peso/estatura, lo calculamos (altura en metros)
    if data.get("imc") is None and data.get("peso") and data.get("estatura"):
        h_m = data["estatura"] / 100.0
        if h_m > 0:
            data["imc"] = round(data["peso"] / (h_m**2), 1)

    return data

# -------- Función para separar las comidas --------
def split_plan_sections(plan_text: str) -> dict:
    """
    Devuelve un dict con claves: DESAYUNO, COMIDA, CENA, MERIENDA.
    Soporta encabezados con numeración, bullets, markdown (** ###), dos puntos, etc.
    """
    sections = {"DESAYUNO": "", "COMIDA": "", "CENA": "", "MERIENDA": ""}

    # Encabezado tolerante:
    # - opcional numeración/bullets/markdown al principio
    # - nombres en singular/plural y variantes (almuerzo/snack)
    header_re = re.compile(
        r"""(?im)            # multiline + ignorecase
        ^\s*                 # espacios iniciales
        (?:[#>*-]+\s*)?      # opcional markdown/bullets (###, -, >)
        (?:\d+\.\s*)?        # opcional "1. "
        (?:\*\*|__)?\s*      # opcional **negrita** / __
        (desayunos?|comidas?|almuerzos?|cenas?|meriendas?|snacks?|merienda\s*o\s*snack)
        \s*(?:\*\*|__)?\s*   # opcional cierre negrita
        :?\s*$               # opcional ":"
        """,
        re.VERBOSE,
    )

    # Encuentra todos los encabezados y sus posiciones
    matches = list(header_re.finditer(plan_text))
    if not matches:
        return sections  # no detectado -> devolver vacío (plantilla mostrará "—")

    # Añade un marcador final al texto para delimitar la última sección
    bounds = [(m.group(1).lower(), m.start(), m.end()) for m in matches]
    bounds.append(("__end__", len(plan_text), len(plan_text)))

    # Recorre por tramos
    for i in range(len(bounds) - 1):
        name, start, end = bounds[i]
        next_start = bounds[i + 1][1]
        body = plan_text[end:next_start].strip()

        # Mapea nombre detectado -> clave destino
        if name.startswith("desayuno"):
            key = "DESAYUNO"
        elif name.startswith(("comida", "almuerzo")):
            key = "COMIDA"
        elif name.startswith("cena"):
            key = "CENA"
        elif name.startswith(("merienda", "snack", "merienda o snack")):
            key = "MERIENDA"
        else:
            continue

        # Acumula texto (por si el modelo repite encabezados)
        sections[key] = (sections[key] + "\n" + body).strip() if sections[key] else body

    return sections

# -------- Función para rellenar el fichero --------
def fill_docx_template(template_file, mapping: dict) -> BytesIO:
    """
    Rellena un .docx sustituyendo cada {{CLAVE}} por su valor en mapping.
    OJO: reasigna el texto de párrafos/celdas (puede perder estilos por-run muy finos).
    """
    doc = Document(template_file)

    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            text = p.text
            for k, v in mapping.items():
                text = text.replace(k, v)
            p.text = text

    def replace_in_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    # reemplaza en el texto de la celda
                    text = cell.text
                    for k, v in mapping.items():
                        text = text.replace(k, v)
                    # reasigna todo el texto de la celda
                    for p in cell.paragraphs:
                        p.text = ""  # limpia
                    cell.paragraphs[0].text = text
                    # también procesa tablas anidadas
                    if cell.tables:
                        replace_in_tables(cell.tables)

    replace_in_paragraphs(doc.paragraphs)
    replace_in_tables(doc.tables)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# -------- UI --------
st.title("🥗 NutriGen - Generador Nutricional Interactivo")

# 1) Subir ficha de Word
uploaded = st.file_uploader("Sube la ficha del cliente (.docx)", type=["docx"])
template_file = st.file_uploader("Sube la PLANTILLA (.docx) que quieres rellenar", type=["docx"])
# 2) Parsear y mostrar un resumen rápido
parsed = {}
if uploaded is not None:
    try:
        parsed = parse_docx(uploaded) or {}
        with st.expander("📄 Datos importados del documento"):
            st.write(parsed)
    except Exception as e:
        st.error(f"No se pudo leer el documento: {e}")

# 3) Form con valores por defecto desde el .docx (si existen)
with st.form("nutri_form"):
    st.subheader("Introduce / revisa los datos")

    nombre = st.text_input("Nombre completo", value=parsed.get("nombre", "") or "")
    edad = st.number_input("Edad", 0, 120, int(parsed.get("edad") or 30))
    sexo = st.selectbox("Sexo", ["Hombre", "Mujer", "Otro"], 
                        index={"hombre":0,"mujer":1}.get(str(parsed.get("sexo") or "").lower(), 2))

    weight_default = float(parsed.get("peso") or 70.0)
    height_default = float(parsed.get("estatura") or 170.0)

    weight = st.number_input("Peso (kg)", 30.0, 300.0, weight_default)
    height = st.number_input("Estatura (cm)", 120.0, 230.0, height_default)

    imc_val = parsed.get("imc")
    st.markdown(f"**IMC (auto):** {imc_val if imc_val is not None else '—'}")

    grasa = st.number_input("% Grasa corporal", 0.0, 70.0, float(parsed.get("grasa") or 0.0))
    masa = st.number_input("% Masa muscular", 0.0, 80.0, float(parsed.get("masa_muscular") or 0.0))

    objetivo = st.text_input("Objetivo nutricional principal", value=parsed.get("objetivo", "") or "")

    st.markdown("---")
    st.subheader("Parámetros del plan")
    preferences = st.text_input("Preferencias alimentarias (ej. mediterránea, vegana...)")
    restrictions = st.text_input("Restricciones (ej. sin gluten, sin lactosa...)")
    calories = st.number_input("Calorías diarias objetivo", 1000, 5000, 2000)
    protein = st.number_input("Proteína diaria objetivo (g)", 20, 300, 120)
    sugar = st.number_input("Azúcar diaria objetivo (g)", 0, 200, 30)
    menu_input = st.text_area("Lista de alimentos disponibles (ej. Arroz = 100, Pollo = 200...)")

    submitted = st.form_submit_button("Generar Plan")

# 4) Generar plan
if submitted:
    user_data = {
        "nombre": nombre,
        "edad": edad,
        "sexo": sexo,
        "peso": weight,
        "estatura": height,
        "imc": imc_val,
        "grasa": grasa,
        "masa_muscular": masa,
        "objetivo": objetivo,
        "preferences": preferences,
        "restrictions": restrictions,
        "calories": calories,
        "protein": protein,
        "sugar": sugar
    }

    st.markdown("### 🧠 Resultado:")
    with st.spinner("Generando el plan..."):
        plan = generate_meal_plan(user_data, menu_input)
        # ✅ GUARDA EN SESIÓN
        st.session_state["user_data"] = user_data
        st.session_state["plan"] = plan
        st.session_state["has_plan"] = True

        st.session_state["raw_plan"] = plan  # si lo usas en otro sitio
        
        secs = split_plan_sections(plan)
        st.session_state["plan_sections"] = secs
        
        st.markdown("---")
        blocks = re.split(r"\n-{3,}\n", plan) if "---" in plan else [plan]
        for block in blocks:
            st.markdown(block)

has_plan = st.session_state.get("has_plan", False)
saved_user_data = st.session_state.get("user_data")
saved_plan = st.session_state.get("plan")

# --- Si hay plantilla subida, la rellenamos y damos descarga ---
if template_file is not None and has_plan and saved_user_data and saved_plan:
    user_data = saved_user_data
    plan = saved_plan
    secs = st.session_state.get("plan_sections", {})

    # macros simple
    try:
        pct_prot = round((user_data["protein"] * 4) / user_data["calories"] * 100)
    except Exception:
        pct_prot = 20
    pct_grasa = 30
    pct_carb = max(0, 100 - pct_prot - pct_grasa)

    reparto_macros = (
        f"Calorías totales: {user_data['calories']} kcal\n"
        f"Proteínas: ~{pct_prot}%\n"
        f"Grasas: ~{pct_grasa}%\n"
        f"Carbohidratos: ~{pct_carb}%"
    )
    recomendaciones = (
        "- Hidratación: 1.5–2 L/día.\n"
        "- Actividad física: 3–5 días/semana, combinando fuerza y cardio.\n"
        "- Sueño: 7–9 horas/noche.\n"
        "- Verduras en 2+ comidas diarias."
    )
    indicaciones = "Adaptar raciones a hambre/saciedad. Priorizar alimentos frescos."
    seguimiento = "Revisión en 4–6 semanas."
    observaciones = "—"


    mapping = {
        "{{NOMBRE_COMPLETO}}": str(user_data.get("nombre") or "—"),
        "{{EDAD}}": str(user_data.get("edad") or "—"),
        "{{SEXO}}": str(user_data.get("sexo") or "—"),
        "{{PESO_KG}}": f"{float(user_data.get('peso') or 0):.1f}" if user_data.get("peso") else "—",
        "{{ESTATURA_CM}}": f"{float(user_data.get('estatura') or 0):.1f}" if user_data.get("estatura") else "—",
        "{{IMC}}": f"{float(user_data.get('imc')):.1f}" if user_data.get("imc") is not None else "—",
        "{{GRASA_PCT}}": f"{float(user_data.get('grasa') or 0):.1f}" if user_data.get("grasa") is not None else "—",
        "{{MASA_MUSCULAR_PCT}}": f"{float(user_data.get('masa_muscular') or 0):.1f}" if user_data.get("masa_muscular") is not None else "—",
        "{{OBJETIVO_PRINCIPAL}}": str(user_data.get("objetivo") or "—"),
        "{{PREFERENCIAS}}": str(user_data.get("preferences") or "—"),
        "{{RESTRICCIONES}}": str(user_data.get("restrictions") or "—"),
        "{{CALORIAS_OBJ}}": str(user_data.get("calories")),
        "{{PROTEINA_OBJ}}": str(user_data.get("protein")),
        "{{AZUCAR_OBJ}}": str(user_data.get("sugar")),
        "{{DESAYUNO}}": secs.get("DESAYUNO", "") or "—",
        "{{COMIDA}}": secs.get("COMIDA", "") or "—",
        "{{CENA}}": secs.get("CENA", "") or "—",
        "{{MERIENDA}}": secs.get("MERIENDA", "") or "—",
        "{{REPARTO_MACROS}}": reparto_macros,
        "{{RECOMENDACIONES}}": recomendaciones,
        "{{INDICACIONES}}": indicaciones,
        "{{SEGUIMIENTO}}": seguimiento,
        "{{OBSERVACIONES}}": observaciones,
    }

    filled = fill_docx_template(template_file, mapping)
    nombre_archivo = f"Plan_{(user_data.get('nombre') or 'cliente').replace(' ', '_')}.docx"

    st.download_button(
        label="⬇️ Descargar plan rellenado (.docx)",
        data=filled,
        file_name=nombre_archivo,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
else:
    st.info("Primero genera el plan y luego podrás descargar la plantilla rellenada.")
